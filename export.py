from docx import Document
from pathlib import Path
from psycopg2.extras import RealDictCursor
from db import get_conn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

TEMPLATE_PATH = "./song_template.docx"
EXPORT_PATH = "./export/songs.docx"

def replace_runs(paragraph, data):
    for r in paragraph.runs:
        for k, v in data.items():
            if k in r.text:
                r.text = r.text.replace(k, v)

def replace_doc(doc, data):
    # параграфы документа
    for p in doc.paragraphs:
        replace_runs(p, data)
    # таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_runs(p, data)

def copy_paragraph(src_paragraph, dst_document):
    p = dst_document.add_paragraph()
    # безопасное выравнивание
    align_map = {
        None: None,
        WD_ALIGN_PARAGRAPH.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    try:
        p.alignment = align_map.get(src_paragraph.alignment, WD_ALIGN_PARAGRAPH.LEFT)
    except Exception:
        p.alignment = None  # fallback

    for r in src_paragraph.runs:
        new_r = p.add_run(r.text)
        new_r.bold = r.bold
        new_r.italic = r.italic
        new_r.underline = r.underline
        new_r.font.name = r.font.name
        new_r.font.size = r.font.size
        if r.font.color.rgb:
            new_r.font.color.rgb = r.font.color.rgb

def copy_table(src_table, dst_document):
    t = dst_document.add_table(rows=0, cols=len(src_table.columns))
    for row in src_table.rows:
        cells = t.add_row().cells
        for i, cell in enumerate(row.cells):
            for p in cell.paragraphs:
                copy_paragraph(p, cells[i].paragraphs[0]._element.getparent().part)

def split_lyrics(lyrics: str):
    fix_keywords = ["Вступление", "Интро", "Тэг", "Интерлют", "Кода", "Проигрыш", "Инструментал"]

    chord_pattern = r'^[A-GH][#b]?m?(?:aug|dim|sus|add)?[0-9]?(?:/[A-GH][#b]?m?)?$'
    chord_token = re.compile(chord_pattern)
    english_letter = re.compile(r'[A-H]')  # любая английская буква аккорда

    special_tokens = set('|:()Xx0123456789')

    def is_chord_line(line: str):
        # если есть английская буква A-H → аккорд
        if english_letter.search(line):
            return True

        tokens = line.strip().split()
        if not tokens:
            return False

        chord_count = 0
        for tok in tokens:
            clean_tok = ''.join(c for c in tok if c not in special_tokens)
            if not clean_tok:
                continue
            if chord_token.match(clean_tok):
                chord_count += 1
        return chord_count / max(len(tokens), 1) > 0.5

    lines = lyrics.splitlines()
    lyrics_text = []

    prev_empty = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if not prev_empty:
                lyrics_text.append("")
            prev_empty = True
            continue
        prev_empty = False

        matched_keyword = next((kw for kw in fix_keywords if stripped.startswith(kw + ":")), None)
        if matched_keyword:
            lyrics_text.append(matched_keyword)
            continue

        if is_chord_line(stripped):
            continue  # убираем строку

        lyrics_text.append(stripped)

    # чистка пустых строк в начале и конце
    while lyrics_text and lyrics_text[0] == "":
        lyrics_text.pop(0)
    while lyrics_text and lyrics_text[-1] == "":
        lyrics_text.pop()

    # нормализация: не более одной пустой строки подряд
    normalized = []
    prev_empty = False
    for line in lyrics_text:
        if line == "":
            if not prev_empty:
                normalized.append(line)
            prev_empty = True
        else:
            normalized.append(line)
            prev_empty = False

    return "", '\n'.join(normalized)

def export():
    Path("./export").mkdir(exist_ok=True)
    with get_conn() as conn:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("""
                SELECT title, title_en, type, key_letter, bpm, has_playback, lyrics, notes
                FROM songs ORDER BY title
            """)
            rows = cur.fetchall()

    if not rows:
        return

    out = Document()

    for i, row in enumerate(rows):
        d = Document(TEMPLATE_PATH)

        # вычисляем chords и lyrics
        raw_lyrics = row.get("lyrics") or ""
        chords, lyrics_only = split_lyrics(raw_lyrics)

        # формируем строку info_line с TYPE_LABEL
        type_label = "Поклонение" if row["type"] == 0 else "Прославление"
        key = row["key_letter"] or ""
        bpm = str(row["bpm"] or "")
        playback_icon = "✓" if row["has_playback"] else "✗"
        info_line = f"{key}    |    ♩ = {bpm}    |    {playback_icon}    |    {type_label}"

        # данные для замены
        data = {
            "{{TITLE}}": row["title"] or "",
            "{{TITLE_EN}}": row["title_en"] or "",
            "{{KEY_LETTER}}  ·  ♩ = {{BPM}}  ·  {{PLAYBACK_ICON}}": info_line,
            "{{LYRICS}}": lyrics_only,
            "{{NOTES}}": row.get("notes") or "",
        }

        replace_doc(d, data)

        # копируем параграфы с run и выравниванием
        for p in d.paragraphs:
            copy_paragraph(p, out)
        # копируем таблицы
        for t in d.tables:
            copy_table(t, out)

        if i < len(rows) - 1:
            out.add_page_break()

    out.save(EXPORT_PATH)

if __name__ == "__main__":
    export()